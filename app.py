import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO

st.set_page_config(page_title="MS WORD Format Checker")
st.title("🧾 MS WORD Format Checker")
st.write("कृपया नीचे विवरण भरें और .docx फाइल अपलोड करें:")

name = st.text_input("👤 छात्र का नाम")
roll = st.text_input("🆔 रोल नंबर")
uploaded_file = st.file_uploader("अपनी Word फ़ाइल अपलोड करें (.docx)", type=["docx"])

def check_formatting(file):
    doc = Document(file)
    score = 0
    feedback = []
    section = doc.sections[0]

    # --- Q1 ---
    if round(section.page_width.inches, 1) == 11.7 and round(section.page_height.inches, 1) == 16.5:
        score += 2.5
        feedback.append("✅ Q1(अ): A3 पेज साइज सेट है")
    else:
        feedback.append("❌ Q1(अ): A3 पेज साइज नहीं है")

    if len(doc.paragraphs) > 1 and 'w:top' in doc.paragraphs[1]._element.xml:
        score += 2.5
        feedback.append("✅ Q1(ब): दूसरे पैराग्राफ में बॉर्डर है")
    else:
        feedback.append("❌ Q1(ब): दूसरे पैराग्राफ में बॉर्डर नहीं है")

    if round(section.top_margin.inches, 1) == 0.6:
        score += 2.5
        feedback.append("✅ Q1(स): टॉप मार्जिन 0.6 इंच है")
    else:
        feedback.append("❌ Q1(स): टॉप मार्जिन 0.6 इंच नहीं है")

    if len(doc.paragraphs) > 3:
        found = any("Define" in run.text and run.bold and run.font.highlight_color for run in doc.paragraphs[3].runs)
        if found:
            score += 2.5
            feedback.append("✅ Q1(द): 'Define' बोल्ड और हाईलाइट किया गया है")
        else:
            feedback.append("❌ Q1(द): 'Define' बोल्ड और हाईलाइट नहीं है")

    # --- Q2 ---
    feedback.append("⚠️ Q2(अ): Watermark की जांच संभव नहीं")
    feedback.append("⚠️ Q2(ब): Header Page No. की जांच संभव नहीं")

    if len(doc.paragraphs) > 1 and round(doc.paragraphs[1].paragraph_format.line_spacing or 0, 2) == 1.5:
        score += 2.5
        feedback.append("✅ Q2(स): पैराग्राफ 2 में लाइन स्पेसिंग 1.5 है")
    else:
        feedback.append("❌ Q2(स): पैराग्राफ 2 में लाइन स्पेसिंग 1.5 नहीं है")

    found_raksha_font = any("रक्षा" in run.text and run.font.name for p in doc.paragraphs for run in p.runs)
    if found_raksha_font:
        score += 2.5
        feedback.append("✅ Q2(द): 'रक्षा' शब्द का फॉन्ट बदला गया है")
    else:
        feedback.append("❌ Q2(द): 'रक्षा' शब्द का फॉन्ट नहीं बदला गया")

    # --- Q3 ---
    underline_removed = all("रक्षा" not in run.text or not run.underline for p in doc.paragraphs for run in p.runs)
    if underline_removed:
        score += 2.5
        feedback.append("✅ Q3(अ): 'रक्षा' का अंडरलाइन हटाया गया है")
    else:
        feedback.append("❌ Q3(अ): 'रक्षा' का अंडरलाइन अभी भी है")

    if len(doc.paragraphs) > 0 and round(doc.paragraphs[0].paragraph_format.line_spacing or 0, 2) == 1.15:
        score += 2.5
        feedback.append("✅ Q3(ब): पैराग्राफ 1 में लाइन स्पेसिंग 1.15 है")
    else:
        feedback.append("❌ Q3(ब): पैराग्राफ 1 में लाइन स्पेसिंग 1.15 नहीं है")

    if len(doc.paragraphs) > 2 and doc.paragraphs[2].style.name.lower().startswith("list"):
        score += 2.5
        feedback.append("✅ Q3(स): पैराग्राफ 3 बुलेट लिस्ट है")
    else:
        feedback.append("❌ Q3(स): पैराग्राफ 3 बुलेट लिस्ट नहीं है")

    feedback.append("⚠️ Q3(द): टेबल की लाइन स्पेसिंग की जांच संभव नहीं")

    # --- Q4 ---
    found_orange = any(run.font.color and run.font.color.rgb and str(run.font.color.rgb) == 'FFA500' for p in doc.paragraphs for run in p.runs)
    if found_orange:
        score += 2.5
        feedback.append("✅ Q4(अ): ऑरेंज फॉन्ट कलर सेट है")
    else:
        feedback.append("❌ Q4(अ): ऑरेंज फॉन्ट कलर सेट नहीं है")

    para = doc.paragraphs[0]
    if para.paragraph_format.line_spacing == Pt(12) and para.paragraph_format.space_after == Pt(18):
        score += 2.5
        feedback.append("✅ Q4(ब): लाइन स्पेसिंग 12pt और आफ्टर 18pt है")
    else:
        feedback.append("❌ Q4(ब): लाइन स्पेसिंग या आफ्टर स्पेसिंग सही नहीं है")

    feedback.append("⚠️ Q4(स): लिंक प्रोसेसिंग की जांच संभव नहीं")
    feedback.append("⚠️ Q4(द): पेज मार्जिन सभी दिशाओं में '12' की जांच सीमित है")

    # --- Q5 ---
    tables = doc.tables
    if tables:
        aligned_right = all(cell.paragraphs[0].alignment == 2 for row in tables[0].rows for cell in row.cells if cell.paragraphs)
        if aligned_right:
            score += 2.5
            feedback.append("✅ Q5(अ): सारणी का डेटा राइट अलाइन्ड है")
        else:
            feedback.append("❌ Q5(अ): सारणी का डेटा राइट अलाइन्ड नहीं है")
    else:
        feedback.append("❌ Q5(अ): कोई सारणी नहीं मिली")

    feedback.append("⚠️ Q5(ब): शैडो इफेक्ट की जांच संभव नहीं")
    feedback.append("⚠️ Q5(स): स्केलिंग इफेक्ट की जांच संभव नहीं")

    return score, feedback

if uploaded_file and name and roll:
    score, results = check_formatting(BytesIO(uploaded_file.read()))
    st.success(f"🎯 {name} (Roll: {roll}) – Total Score: {score}/50")
    st.write("### 📋 फीडबैक:")
    for r in results:
        st.write(r)

    report = f"नाम: {name}\\nरोल: {roll}\\nअंक: {score}/50\\n\\n" + "\\n".join(results)
    st.download_button("📥 रिपोर्ट डाउनलोड करें", report, file_name=f"{roll}_report.txt", mime="text/plain")
elif uploaded_file and (not name or not roll):
    st.warning("कृपया नाम और रोल नंबर भरें")
